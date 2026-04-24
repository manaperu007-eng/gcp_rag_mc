##############################################################################
# backend/ingestion/processors/document_ai_processor.py
# Extract text from PDF/DOCX via Google Cloud Document AI
# with local fallbacks (pypdf / python-docx) when processor ID is unset.
##############################################################################
from __future__ import annotations

import logging
import os
from typing import Any, Dict, List, Optional

from shared.config import Settings

logger = logging.getLogger(__name__)


class DocumentAIProcessor:
    """
    Extracts structured page text from PDF, DOCX, and DOC files.

    Strategy (in order of preference):
      1. Google Cloud Document AI  — if ``document_ai_processor_id`` is set.
      2. Local pypdf               — for PDF files.
      3. Local python-docx         — for DOCX/DOC files.

    All paths return a list of page dicts::

        [{"text": "...", "page_number": 1, "section_title": None}, ...]
    """

    def __init__(self, settings: Settings) -> None:
        self.settings = settings
        self._doc_ai_client: Optional[Any] = None

    # ── Public API ────────────────────────────────────────────────────────

    def extract_text(self, file_path: str, file_type: str) -> List[Dict[str, Any]]:
        """
        Extract text pages from *file_path*.

        :param file_path: Local path to the downloaded file.
        :param file_type: ``pdf`` | ``docx`` | ``doc``
        :returns: List of page dicts.
        """
        processor_id = self.settings.document_ai_processor_id
        use_doc_ai = bool(processor_id) and file_type in ("pdf", "docx", "doc")

        if use_doc_ai:
            try:
                return self._extract_via_document_ai(file_path, file_type)
            except Exception as exc:
                logger.warning(
                    "Document AI extraction failed (%s), falling back to local parser", exc
                )

        # Local fallbacks
        if file_type == "pdf":
            return self._extract_pdf_local(file_path)
        if file_type in ("docx", "doc"):
            return self._extract_docx_local(file_path)

        raise ValueError(f"Unsupported file type for DocumentAIProcessor: {file_type}")

    # ── Google Cloud Document AI ──────────────────────────────────────────

    def _extract_via_document_ai(
        self, file_path: str, file_type: str
    ) -> List[Dict[str, Any]]:
        """Call Document AI Online Processing API."""
        from google.cloud import documentai  # noqa: PLC0415

        mime_map = {
            "pdf":  "application/pdf",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "doc":  "application/msword",
        }
        mime_type = mime_map.get(file_type, "application/octet-stream")

        with open(file_path, "rb") as f:
            content = f.read()

        client = self._get_doc_ai_client()
        processor_name = client.processor_path(
            self.settings.project_id,
            self.settings.document_ai_location,
            self.settings.document_ai_processor_id,
        )

        raw_document = documentai.RawDocument(content=content, mime_type=mime_type)
        request = documentai.ProcessRequest(
            name=processor_name,
            raw_document=raw_document,
        )

        result = client.process_document(request=request)
        document = result.document

        pages: List[Dict[str, Any]] = []
        for i, page in enumerate(document.pages, start=1):
            # Collect all paragraph texts for this page
            page_texts: List[str] = []
            for block in page.blocks:
                text = _layout_text(block.layout, document.text)
                if text.strip():
                    page_texts.append(text.strip())

            pages.append({
                "text": "\n".join(page_texts),
                "page_number": i,
                "section_title": None,
            })

        logger.info(
            "Document AI extracted %d pages from %s", len(pages), os.path.basename(file_path)
        )
        return pages

    def _get_doc_ai_client(self):
        if self._doc_ai_client is None:
            from google.cloud import documentai  # noqa: PLC0415
            opts = {"api_endpoint": f"{self.settings.document_ai_location}-documentai.googleapis.com"}
            self._doc_ai_client = documentai.DocumentProcessorServiceClient(client_options=opts)
        return self._doc_ai_client

    # ── Local PDF fallback ────────────────────────────────────────────────

    @staticmethod
    def _extract_pdf_local(file_path: str) -> List[Dict[str, Any]]:
        """Extract text page-by-page using pypdf."""
        try:
            import pypdf  # noqa: PLC0415
        except ImportError:
            try:
                import PyPDF2 as pypdf  # noqa: PLC0415, N814
            except ImportError as exc:
                raise ImportError(
                    "pypdf (or PyPDF2) is required for local PDF parsing. "
                    "Install with: pip install pypdf"
                ) from exc

        pages: List[Dict[str, Any]] = []
        with open(file_path, "rb") as f:
            reader = pypdf.PdfReader(f)
            for i, page in enumerate(reader.pages, start=1):
                try:
                    text = page.extract_text() or ""
                except Exception:
                    text = ""
                pages.append({
                    "text": text.strip(),
                    "page_number": i,
                    "section_title": None,
                })

        logger.info(
            "pypdf extracted %d pages from %s", len(pages), os.path.basename(file_path)
        )
        return pages

    # ── Local DOCX fallback ───────────────────────────────────────────────

    @staticmethod
    def _extract_docx_local(file_path: str) -> List[Dict[str, Any]]:
        """
        Extract text from DOCX/DOC using python-docx.
        DOCX does not have true pages, so we split on Heading styles to simulate pages.
        """
        try:
            import docx  # noqa: PLC0415
        except ImportError as exc:
            raise ImportError(
                "python-docx is required for DOCX parsing. "
                "Install with: pip install python-docx"
            ) from exc

        document = docx.Document(file_path)

        pages: List[Dict[str, Any]] = []
        current_section: List[str] = []
        current_heading: Optional[str] = None
        page_num = 1

        def _flush():
            nonlocal page_num, current_section, current_heading
            text = "\n".join(current_section).strip()
            if text:
                pages.append({
                    "text": text,
                    "page_number": page_num,
                    "section_title": current_heading,
                })
                page_num += 1
            current_section = []

        for para in document.paragraphs:
            style = para.style.name if para.style else ""
            if "Heading" in style:
                _flush()
                current_heading = para.text.strip() or None
            else:
                if para.text.strip():
                    current_section.append(para.text.strip())

        _flush()  # Flush final section

        if not pages:
            # Fallback: treat entire document as one page
            full_text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
            pages = [{"text": full_text, "page_number": 1, "section_title": None}]

        logger.info(
            "python-docx extracted %d sections from %s",
            len(pages),
            os.path.basename(file_path),
        )
        return pages


# ── Helpers ───────────────────────────────────────────────────────────────────

def _layout_text(layout, full_text: str) -> str:
    """Extract the text for a Document AI Layout from the full document text."""
    response = ""
    for segment in layout.text_anchor.text_segments:
        start = int(segment.start_index)
        end = int(segment.end_index)
        response += full_text[start:end]
    return response
